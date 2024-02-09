import streamlit as st
from validators import email
import re
import google.generativeai as genai
import base64
import html
from docx import Document
from io import BytesIO

class ResumeMaker:
    def __init__(self):
        self.resume_content = "<html><head><style>body {font-family: Arial, sans-serif;}</style></head><body>"
        
    def add_heading(self, text, level=1):
        self.resume_content += f"<h{level} style='color: #333333;'>{text}</h{level}>"
    
    def add_paragraph(self, text):
        self.resume_content += f"<p>{text}</p>"
        
    def generation_config(self , prompt_parts):
        genai.configure(api_key="AIzaSyDiLrqZQtcnX_jPDMvJd6VVju29lDXw8GM")  # Replace with your actual API key

        generation_config = {
            "temperature": 0.9,
            "top_p": 1,
            "top_k": 1,
            "max_output_tokens": 2048,
        }

        safety_settings = [
            {
                "category": "HARM_CATEGORY_HARASSMENT",
                "threshold": "BLOCK_MEDIUM_AND_ABOVE"
            },
            {
                "category": "HARM_CATEGORY_HATE_SPEECH",
                "threshold": "BLOCK_MEDIUM_AND_ABOVE"
            },
            {
                "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
                "threshold": "BLOCK_MEDIUM_AND_ABOVE"
            },
            {
                "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
                "threshold": "BLOCK_MEDIUM_AND_ABOVE"
            }
        ]

        model = genai.GenerativeModel(model_name="gemini-pro",
                                      generation_config=generation_config,
                                      safety_settings=safety_settings)

        response = model.generate_content(prompt_parts)
        return response.text
    
    def save_resume(self):
        self.resume_content += "</body></html>"
        return self.resume_content


def get_word_document_download_link(content, file_label='File', file_type='docx'):
    """
    Generates a downloadable link for a Word document.

    Args:
        content (str): The content of the Word document.
        file_label (str, optional): The label to display for the file. Defaults to 'File'.
        file_type (str, optional): The file type (e.g., 'docx'). Defaults to 'docx'.

    Returns:
        str: The HTML code for a downloadable link.
    """
    try:
        # Create a BytesIO object to store the document
        f = BytesIO()
        
        # Create a new Word document
        doc = Document()
        
        # Add content to the Word document
        doc.add_paragraph(content)
        
        # Save the document to the BytesIO object
        doc.save(f)
        
        # Reset the BytesIO object's position to the beginning
        f.seek(0)
        
        # Convert the BytesIO object to bytes
        doc_bytes = f.getvalue()
        
        # Base64 encode and escape special characters for safe inclusion in HTML
        b64_encoded_content = base64.b64encode(doc_bytes).decode()
        escaped_content = html.escape(b64_encoded_content)
        
        # Create the download link using the correct content type and filename
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{escaped_content}" download="{file_label}.{file_type}">Click here to download {file_label}</a>'

        return href

    except Exception as e:
        # Handle exceptions gracefully with informative messages
        error_message = f"An error occurred while creating the download link: {e}"
        st.error(error_message)
        return None


def validate_email(email):
    email_regex = r"^[a-zA-Z0-9.!#$%&'*+/=?^_`{|}~-]+@[a-zA-Z0-9-]+(?:\.[a-zA-Z0-9-]+)*$"
    return bool(re.match(email_regex, email))


# Define your ResumeMaker class and other functions here

st.title("FJ Resume Maker")

name = st.text_input("Please enter your full name:", key="full_name")
email_address = st.text_input("Please enter your email address:", key="email_address")
phone = st.text_input("Please enter your phone number (optional):", key="phone_number")
education_prompt = st.text_input("Enter Your Education:", key="education")
work_prompt = st.text_input("Enter Your Work Experience:", key="work_experience")

resume_content = None  # Initialize the variable outside the if block

if st.button("Create Resume"):
    if not name or not email_address:
        st.warning("Name and email address are required.")
    elif not validate_email(email_address):
        st.warning("Invalid email format. Please enter a valid email address.")
    else:
        with st.spinner("Generating your resume..."):
            resume_maker = ResumeMaker()
            resume_maker.add_heading(name, level=1)
            resume_maker.add_paragraph(f"Contact Information:")
            resume_maker.add_paragraph(f"* Email: {email_address}")
            resume_maker.add_paragraph(f"* Phone: {phone}")

            education_details = resume_maker.generation_config(education_prompt)
            resume_maker.add_heading("Education", level=2)
            resume_maker.add_paragraph(education_details.strip())

            work_details = resume_maker.generation_config(work_prompt)
            resume_maker.add_heading("Work Experience", level=3)
            resume_maker.add_paragraph(work_details.strip())

            resume_content = resume_maker.save_resume()  # Store the generated resume content

            # Display editable resume content in a text box
            st.subheader("Generated Resume Content")
            st.markdown(resume_content, unsafe_allow_html=True)  # Render HTML content using markdown

            if st.button("Download Resume"):
                st.markdown(get_word_document_download_link(resume_content, 'Edited_Resume', 'docx'), unsafe_allow_html=True)


